#!/usr/bin/env python3
"""
TXT/MD/DOCX 文件同步脚本
功能：扫描源目录，转换为静态文件，生成索引
"""

import sys
import os
import re
import json
import shutil
from pathlib import Path
from datetime import datetime
from fnmatch import fnmatch, fnmatchcase

SKIP_NAMES = ['.git', '__pycache__', 'node_modules', '.github', 'reader', 'scripts']
ALLOWED_EXTENSIONS = {'.txt', '.md', '.docx'}

CONFIG_FILE = Path(__file__).parent.parent / 'reader' / 'config.json'

is_github_actions = os.environ.get('GITHUB_ACTIONS') == 'true' or os.environ.get('CI') == 'true'

if is_github_actions:
    current_dir = Path.cwd()
    if 'tools' in str(current_dir).lower():
        CONFIG_FILE = current_dir / 'reader' / 'config.json'

def log_info(msg, module='Sync'):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f'[{module}][{timestamp}] {msg}')

def log_error(msg, module='Sync'):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f'[{module}][{timestamp}] [ERROR] {msg}')

_config_cache = None

def load_config():
    global _config_cache
    if _config_cache is not None:
        return _config_cache
    
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
            log_info(f'配置文件加载成功: {CONFIG_FILE}')
            _config_cache = config
            return config
        except (json.JSONDecodeError, IOError) as e:
            log_error(f'配置文件加载错误: {e}')
            return {}
    else:
        log_error(f'配置文件不存在: {CONFIG_FILE}')
        return {}

def get_source_dir():
    return load_config().get('source_dir', '')

def get_docs_dir():
    sync_root_dir = os.environ.get('SYNC_ROOT_DIR')
    if sync_root_dir:
        root_dir = Path(sync_root_dir)
    else:
        root_dir = Path(__file__).parent.parent
    return root_dir / 'reader' / 'docs'

def get_exclude_patterns():
    return load_config().get('exclude_patterns', [])

def get_exclude_files():
    return load_config().get('exclude_files', [])

def get_site_title():
    return load_config().get('site_title', '文档阅读器')

def get_sidebar_title():
    return load_config().get('sidebar_title', '文档目录')

def get_enable_search():
    return load_config().get('enable_search', True)

def get_home_page():
    return load_config().get('home_page', '')

def get_github_repo():
    return load_config().get('github_repo', '')

def should_skip(path):
    return any(path.name.startswith(s) for s in SKIP_NAMES)

def match_patterns(rel_path, patterns):
    rel_str = str(rel_path).replace('\\', '/')
    for pattern in patterns:
        if pattern.startswith('**/') or fnmatch(rel_str, pattern):
            return True
        if fnmatch(rel_str, pattern.replace('*', '**')):
            return True
    return False

def should_exclude(file_path, source_dir):
    rel_path = file_path.relative_to(source_dir)
    if match_patterns(rel_path, get_exclude_patterns()):
        return True
    if rel_path.name in get_exclude_files():
        return True
    if file_path.suffix not in ALLOWED_EXTENSIONS:
        log_info(f'跳过(不支持类型): {rel_path}', 'Sync-Filter')
        return True
    return False

def escape_html(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

def process_inline(text):
    text = escape_html(text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', text)
    text = re.sub(r'\`([^`]+)\`', r'<code>\1</code>', text)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2" target="_blank">\1</a>', text)
    text = re.sub(r'(?<![\("])(https?://[^\s<">]+)', r'<a href="\1" target="_blank">\1</a>', text)
    text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1">', text)
    return text

def convert_markdown(text):
    html = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        trimmed = line.strip()
        
        if not trimmed:
            html.append('<br>')
            i += 1
            continue
        
        if trimmed.startswith('# '):
            html.append(f'<h1>{process_inline(trimmed[2:])}</h1>')
            i += 1
        elif trimmed.startswith('## '):
            html.append(f'<h2>{process_inline(trimmed[3:])}</h2>')
            i += 1
        elif trimmed.startswith('### '):
            html.append(f'<h3>{process_inline(trimmed[4:])}</h3>')
            i += 1
        elif trimmed == '---':
            html.append('<hr>')
            i += 1
        elif trimmed.startswith('> '):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('> '):
                quote_lines.append(process_inline(lines[i].strip()[2:]))
                i += 1
            html.append(f'<blockquote>{"<br>".join(quote_lines)}</blockquote>')
        elif trimmed.startswith('```'):
            lang = trimmed[3:].strip() if len(trimmed) > 3 else ''
            code_start = i + 1
            code_end = i + 1
            while code_end < len(lines) and not lines[code_end].rstrip().startswith('```'):
                code_end += 1
            
            if lang == 'mermaid' or lang.startswith('gantt') or lang.startswith('flowchart') or lang.startswith('graph'):
                content_lines = lines[code_start:code_end]
                if lang == 'mermaid':
                    mermaid_code = '\n'.join(content_lines)
                else:
                    mermaid_code = lang + '\n' + '\n'.join(content_lines)
                mermaid_code = mermaid_code.strip()
                html.append(f'<div class="mermaid">{mermaid_code}</div>')
            else:
                code_lines = [escape_html(line) for line in lines[code_start:code_end]]
                html.append(f'<pre><code class="language-{lang}">{"<br>".join(code_lines)}</code></pre>')
            
            i = code_end + 1 if code_end < len(lines) else code_end
        elif trimmed.startswith('- '):
            list_items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                list_items.append(f'<li>{process_inline(lines[i].strip()[2:])}</li>')
                i += 1
            html.append(f'<ul>{"".join(list_items)}</ul>')
        elif re.match(r'^\d+\.\s', trimmed):
            list_items = []
            ordered_pattern = r'^\d+\.\s*'
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                list_items.append(f'<li>{process_inline(re.sub(ordered_pattern, "", lines[i].strip()))}</li>')
                i += 1
            html.append(f'<ol>{"".join(list_items)}</ol>')
        elif trimmed.startswith('|') and '|' in trimmed:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            html.append(render_table(table_lines))
        else:
            html.append(f'<p>{process_inline(trimmed)}</p>')
            i += 1
    
    return '\n'.join(html)

def render_table(table_lines):
    if len(table_lines) < 2:
        return ''.join(f'<p>{escape_html(line)}</p>' for line in table_lines)
    
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    
    if len(rows) < 2:
        return ''.join(f'<p>{escape_html(line)}</p>' for line in table_lines)
    
    header_row = rows[0]
    align = []
    for cell in header_row:
        if cell.endswith(':'):
            align.append('style="text-align:left"')
        elif cell.startswith(':') and cell.endswith(':'):
            align.append('style="text-align:center"')
        elif cell.startswith(':'):
            align.append('style="text-align:right"')
        else:
            align.append('')
    
    html = ['<table class="md-table">']
    
    html.append('<thead><tr>')
    for j, cell in enumerate(header_row):
        style = align[j] if j < len(align) else ''
        html.append(f'<th {style}>{escape_html(cell)}</th>')
    html.append('</tr></thead>')
    
    if len(rows) > 1 and not all(c.startswith('-') or c.startswith(':') for c in rows[1]):
        body_start = 1
    else:
        sep_idx = None
        for idx in range(1, len(rows)):
            row_text = '|'.join(rows[idx])
            if re.match(r'^[\s|:,\-\d]+$', row_text):
                sep_idx = idx
                break
        body_start = (sep_idx + 1) if sep_idx else 2
    
    if body_start < len(rows):
        html.append('<tbody>')
        for r in range(body_start, len(rows)):
            html.append('<tr>')
            for j, cell in enumerate(rows[r]):
                style = align[j] if j < len(align) else ''
                html.append(f'<td {style}>{escape_html(cell)}</td>')
            html.append('</tr>')
        html.append('</tbody>')
    
    html.append('</table>')
    return '\n'.join(html)

def convert_docx(docx_path):
    try:
        from docx import Document
        doc = Document(docx_path)
        
        all_sizes = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    all_sizes.append(run.font.size.pt)
        
        if all_sizes:
            min_size = min(all_sizes)
            max_size = max(all_sizes)
        else:
            min_size, max_size = 12, 24
        
        def get_level(size):
            avg_size = sum(all_sizes) / len(all_sizes) if all_sizes else 12
            if size > avg_size + 3.0:
                return 1
            elif size > avg_size + 2.0:
                return 2
            elif size > avg_size + 1.0:
                return 3
            else:
                return 0
        
        html = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                html.append('<br>')
                continue
            
            para_size = None
            for run in para.runs:
                if run.font.size:
                    para_size = run.font.size.pt
                    break
            
            if para_size:
                level = get_level(para_size)
                if level == 1:
                    html.append(f'<h1>{escape_html(text)}</h1>')
                elif level == 2:
                    html.append(f'<h2>{escape_html(text)}</h2>')
                elif level == 3:
                    html.append(f'<h3>{escape_html(text)}</h3>')
                else:
                    html.append(f'<p>{escape_html(text)}</p>')
            else:
                if text.startswith('-'):
                    html.append(f'<li>{escape_html(text[1:].strip())}</li>')
                else:
                    html.append(f'<p>{escape_html(text)}</p>')
        
        return '\n'.join(html)
    except ImportError:
        log_error('请安装 python-docx 库: pip install python-docx', 'Sync-DOCX')
        raise ImportError('请安装 python-docx 库: pip install python-docx')
    except Exception as e:
        log_error(f'Word 文档转换失败: {e}', 'Sync-DOCX')
        raise

def render_content(content, file_path):
    try:
        if file_path.suffix == '.md':
            return convert_markdown(content)
        elif file_path.suffix == '.docx':
            return convert_docx(file_path)
        else:
            return convert_markdown(content)
    except Exception as e:
        log_error(f'内容渲染失败: {file_path} - {e}', 'Sync-Render')
        raise

def scan_directory(source_dir, root_dir=None):
    if root_dir is None:
        root_dir = source_dir
    items = []
    source_path = Path(source_dir)
    if not source_path.exists():
        return []
    
    for path in sorted(source_path.iterdir(), key=lambda x: (x.is_file(), x.name)):
        if should_skip(path):
            continue
        if path.is_dir():
            children = scan_directory(path, root_dir)
            if children:
                items.append({'type': 'folder', 'name': path.name, 'children': children})
        elif path.is_file():
            if should_exclude(path, root_dir):
                log_info(f'排除: {path.relative_to(root_dir)}', 'Sync-Scan')
                continue
            if path.suffix in ['.txt', '.md', '.docx']:
                rel_path = path.relative_to(root_dir)
                ext = path.suffix
                output_path = str(rel_path).replace('\\', '/')
                if ext != '.txt':
                    output_path = str(rel_path.with_suffix('.html')).replace('\\', '/')
                items.append({
                    'type': 'file',
                    'name': path.name,
                    'path': output_path,
                    'title': path.stem.replace('-', ' ').replace('_', ' ')
                })
    return items

def cleanup_orphaned_files(source_dir, dest_dir):
    source_path = Path(source_dir)
    dest_path = Path(dest_dir)
    
    if not dest_path.exists():
        return
    
    valid_files = set()
    for path in source_path.rglob('*.txt'):
        rel_path = path.relative_to(source_path)
        valid_files.add(str(rel_path))
    
    for path in source_path.rglob('*.md'):
        rel_path = path.relative_to(source_path)
        valid_files.add(str(rel_path.with_suffix('.html')))
    
    for path in source_path.rglob('*.docx'):
        rel_path = path.relative_to(source_path)
        valid_files.add(str(rel_path.with_suffix('.html')))
    
    deleted_count = 0
    for html_file in dest_path.rglob('*.html'):
        rel_path = html_file.relative_to(dest_path)
        rel_str = str(rel_path).replace('\\', '/')
        if rel_str not in valid_files:
            try:
                html_file.unlink()
                log_info(f'删除: {rel_str}', 'Sync-Cleanup')
                deleted_count += 1
            except Exception as e:
                log_error(f'删除文件失败: {rel_str} - {e}', 'Sync-Cleanup')
    
    for txt_file in dest_path.rglob('*.txt'):
        rel_path = txt_file.relative_to(dest_path)
        rel_str = str(rel_path).replace('\\', '/')
        if rel_str not in valid_files:
            try:
                txt_file.unlink()
                log_info(f'删除: {rel_str}', 'Sync-Cleanup')
                deleted_count += 1
            except Exception as e:
                log_error(f'删除文件失败: {rel_str} - {e}', 'Sync-Cleanup')
    
    def remove_empty_dirs(dir_path):
        for child in sorted(dir_path.iterdir(), reverse=True):
            if child.is_dir():
                remove_empty_dirs(child)
                try:
                    child.rmdir()
                    log_info(f'删除空目录: {child.relative_to(dest_path)}', 'Sync-Cleanup')
                except OSError:
                    pass
    
    remove_empty_dirs(dest_path)
    
    if deleted_count > 0:
        log_info(f'清理完成，共删除 {deleted_count} 个文件', 'Sync-Cleanup')

def copy_and_convert_files(source_dir, dest_dir):
    dest_dir.mkdir(parents=True, exist_ok=True)
    source_path = Path(source_dir)
    
    txt_count = 0
    md_count = 0
    docx_count = 0
    error_count = 0
    
    for path in source_path.rglob('*.txt'):
        if should_exclude(path, source_path):
            continue
        rel_path = path.relative_to(source_path)
        dest_path = dest_dir / rel_path
        dest_path.parent.mkdir(parents=True, exist_ok=True)
        if dest_path.exists() and dest_path.stat().st_mtime >= path.stat().st_mtime:
            log_info(f'跳过(未修改): {rel_path}', 'Sync-TXT')
        else:
            try:
                shutil.copy2(path, dest_path)
                log_info(f'复制: {rel_path}', 'Sync-TXT')
                txt_count += 1
            except Exception as e:
                log_error(f'复制失败: {rel_path} - {e}', 'Sync-TXT')
                error_count += 1
    
    for path in source_path.rglob('*.md'):
        if should_exclude(path, source_path):
            continue
        rel_path = path.relative_to(source_path)
        dest_path = dest_dir / rel_path.with_suffix('.html')
        dest_path.parent.mkdir(parents=True, exist_ok=True)
        if dest_path.exists() and dest_path.stat().st_mtime >= path.stat().st_mtime:
            log_info(f'跳过(未修改): {rel_path}', 'Sync-MD')
        else:
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
                html = render_content(content, path)
                with open(dest_path, 'w', encoding='utf-8') as f:
                    f.write(html)
                log_info(f'转换: {rel_path} → {rel_path.with_suffix(".html")}', 'Sync-MD')
                md_count += 1
            except Exception as e:
                if dest_path.exists():
                    dest_path.unlink()
                log_error(f'转换失败: {rel_path} - {e}', 'Sync-MD')
                error_count += 1
    
    for path in source_path.rglob('*.docx'):
        if should_exclude(path, source_path):
            continue
        rel_path = path.relative_to(source_path)
        dest_path = dest_dir / rel_path.with_suffix('.html')
        dest_path.parent.mkdir(parents=True, exist_ok=True)
        if dest_path.exists() and dest_path.stat().st_mtime >= path.stat().st_mtime:
            log_info(f'跳过(未修改): {rel_path}', 'Sync-DOCX')
        else:
            try:
                html = render_content('', path)
                with open(dest_path, 'w', encoding='utf-8') as f:
                    f.write(html)
                log_info(f'转换: {rel_path} → {rel_path.with_suffix(".html")}', 'Sync-DOCX')
                docx_count += 1
            except Exception as e:
                if dest_path.exists():
                    dest_path.unlink()
                log_error(f'转换失败: {rel_path} - {e}', 'Sync-DOCX')
                error_count += 1
    
    log_info(f'文件处理完成: TXT={txt_count}, MD={md_count}, DOCX={docx_count}, ERROR={error_count}', 'Sync-Core')

def generate_index(items, output_file):
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(items, f, ensure_ascii=False, indent=2)
        log_info(f'生成索引: {output_file}', 'Sync-Index')
    except Exception as e:
        log_error(f'生成索引失败: {e}', 'Sync-Index')

def main():
    log_info('=' * 60, 'Sync')
    log_info('开始同步', 'Sync')
    log_info('=' * 60, 'Sync')
    
    sync_root_dir = os.environ.get('SYNC_ROOT_DIR')
    if sync_root_dir:
        root_dir = Path(sync_root_dir)
    else:
        root_dir = Path(__file__).parent.parent
    
    source_dir_name = get_source_dir()
    if not source_dir_name:
        txt_dir = root_dir / 'txt'
        if txt_dir.exists() and txt_dir.is_dir():
            source_dir_name = 'txt'
            log_info(f'自动检测到源目录: txt/', 'Sync')
    
    if source_dir_name:
        source_dir = root_dir / source_dir_name
    else:
        source_dir = root_dir
    
    docs_dir = root_dir / 'reader' / 'docs'
    index_file = root_dir / 'reader' / 'index.json'
    github_repo = get_github_repo()
    
    log_info(f'根目录: {root_dir}', 'Sync')
    log_info(f'源目录: {source_dir.relative_to(root_dir) if source_dir != root_dir else "/"}', 'Sync')
    log_info(f'目标目录: {docs_dir}', 'Sync')
    log_info(f'索引文件: {index_file}', 'Sync')
    log_info(f'GitHub 仓库: {github_repo}', 'Sync')
    log_info(f'时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 'Sync')
    
    if not source_dir.exists():
        log_error(f'源目录不存在: {source_dir}', 'Sync')
        return
    
    log_info('清理已删除的文件...', 'Sync')
    cleanup_orphaned_files(source_dir, docs_dir)
    
    items = []
    if source_dir.exists():
        log_info('扫描目录结构...', 'Sync')
        items = [{
            'type': 'folder',
            'name': source_dir_name,
            'children': scan_directory(source_dir, source_dir)
        }]
        log_info(f'扫描完成，共 {len(items[0]["children"])} 个项目', 'Sync')
    
    log_info('复制和转换文件...', 'Sync')
    copy_and_convert_files(source_dir, docs_dir)
    
    log_info('生成索引文件...', 'Sync')
    generate_index(items, index_file)
    
    log_info('=' * 60, 'Sync')
    log_info('同步完成！', 'Sync')
    log_info('=' * 60, 'Sync')

if __name__ == '__main__':
    main()
